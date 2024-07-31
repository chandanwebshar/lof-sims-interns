import streamlit as st
import markdown2
import json
import requests
import fitz  # PyMuPDF for reading PDF files
from sim_prompts import *  # Ensure this import provides the needed functionality
from bs4 import BeautifulSoup
from fpdf import FPDF
from docx import Document
from sqlalchemy import create_engine, Column, Integer, String, Text, MetaData, Index
from sqlalchemy.orm import sessionmaker, declarative_base
import pandas as pd

# Database setup
DATABASE_URL = "sqlite:///app_data.db"  # SQLite database

engine = create_engine(DATABASE_URL)
Session = sessionmaker(bind=engine)
session = Session()
Base = declarative_base()

# Define the models
class Transcript(Base):
    __tablename__ = 'transcripts'
    id = Column(Integer, primary_key=True, autoincrement=True)
    content = Column(Text, nullable=False)
    role = Column(String, nullable=False)
    specialty = Column(String, nullable=True)
    __table_args__ = (Index('transcript_content_idx', 'content'),)

class Assessment(Base):
    __tablename__ = 'assessments'
    id = Column(Integer, primary_key=True, autoincrement=True)
    content = Column(Text, nullable=False)
    role = Column(String, nullable=False)
    specialty = Column(String, nullable=True)
    __table_args__ = (Index('assessment_content_idx', 'content'),)

class CaseDetails(Base):
    __tablename__ = 'case_details'
    id = Column(Integer, primary_key=True, autoincrement=True)
    saved_name = Column(String, nullable=False)
    content = Column(Text, nullable=False)
    role = Column(String, nullable=False)
    specialty = Column(String, nullable=True)
    __table_args__ = (Index('case_details_content_idx', 'content'),)

class LabResults(Base):
    __tablename__ = 'lab_results'
    id = Column(Integer, primary_key=True, autoincrement=True)
    saved_name = Column(String, nullable=False)
    content = Column(Text, nullable=False)
    __table_args__ = (Index('lab_results_content_idx', 'content'),)

# Drop existing lab_results table if it exists
meta = MetaData()
meta.reflect(bind=engine)
if 'lab_results' in meta.tables:
    lab_results_table = meta.tables['lab_results']
    try:
        lab_results_table.drop(engine)
    except Exception as e:
        st.write(f"Error dropping table lab_results: {e}")

# Create tables
Base.metadata.create_all(engine)

st.set_page_config(
    page_title='Simulated Case Generator',
    page_icon='🌌',
    layout="wide",
    initial_sidebar_state='auto'
)
with st.sidebar:
    st.image("static/er_bays.jpeg", use_column_width=True)

class PDF(FPDF):
    def header(self):
        self.set_font("Arial", "B", 12)
        self.cell(0, 10, self.title, 0, 1, "C")
        self.ln(10)

    def chapter_title(self, title, level=1):
        if level == 1:
            self.set_font("Arial", "B", 16)
            self.ln(10)
        elif level == 2:
            self.set_font("Arial", "B", 14)
            self.ln(8)
        elif level == 3:
            self.set_font("Arial", "B", 12)
            self.ln(6)
        self.cell(0, 10, title, 0, 1, "L")
        self.ln(2)

    def chapter_body(self, body):
        self.set_font("Arial", "", 12)
        self.multi_cell(0, 10, body)
        self.ln()

    def add_list(self, items, is_ordered=False):
        self.set_font("Arial", "", 12)
        for i, item in enumerate(items, start=1):
            text = f"{i}. {item}" if is_ordered else f"- {item}"
            self.multi_cell(0, 10, text)
        self.ln()

def html_to_pdf(html_content, name):
    # Use BeautifulSoup to parse the HTML
    soup = BeautifulSoup(html_content, "html.parser")
    
    # Extract title for the document
    case_title_tag = soup.find("h1")
    case_title = case_title_tag.get_text() if case_title_tag else "Document"
    
    # Create PDF instance with dynamic title
    pdf = PDF()
    pdf.title = case_title
    pdf.add_page()

    # Process each section of the HTML
    for element in soup.find_all(["h1", "h2", "h3", "p", "ul", "ol", "li", "hr"]):
        if element.name == "h1":
            pdf.chapter_title(element.get_text(), level=1)
        elif element.name == "h2":
            if "Patient Door Chart" in element.get_text():
                pdf.add_page()
            pdf.chapter_title(element.get_text(), level=2)
        elif element.name == "h3":
            pdf.chapter_title(element.get_text(), level=3)
        elif element.name == "p":
            pdf.chapter_body(element.get_text())
        elif element.name == "ul":
            items = [li.get_text() for li in element.find_all("li")]
            pdf.add_list(items, is_ordered=False)
        elif element.name == "ol":
            items = [li.get_text() for li in element.find_all("li")]
            pdf.add_list(items, is_ordered=True)
        elif element.name == "hr":
            pdf.add_page()
    
    # Output the PDF
    pdf.output(name)

def html_to_docx(html_content, name):
    # Use BeautifulSoup to parse the HTML
    soup = BeautifulSoup(html_content, "html.parser")
    
    # Create DOCX instance
    doc = Document()

    # Extract title for the document
    case_title_tag = soup.find("h1")
    case_title = case_title_tag.get_text() if case_title_tag else "Document"
    doc.add_heading(case_title, level=1)
    
    # Process each section of the HTML
    for element in soup.find_all(["h1", "h2", "h3", "p", "ul", "ol", "li", "hr"]):
        if element.name == "h1":
            doc.add_heading(element.get_text(), level=1)
        elif element.name == "h2":
            doc.add_heading(element.get_text(), level=2)
        elif element.name == "h3":
            doc.add_heading(element.get_text(), level=3)
        elif element.name == "p":
            doc.add_paragraph(element.get_text())
        elif element.name == "ul":
            for li in element.find_all("li"):
                doc.add_paragraph(f"- {li.get_text()}")
        elif element.name == "ol":
            for i, li in enumerate(element.find_all("li"), start=1):
                doc.add_paragraph(f"{i}. {li.get_text()}")
        elif element.name == "hr":
            doc.add_page_break()
    
    # Output the DOCX
    doc.save(name)

def init_session():
    if "final_case" not in st.session_state:
        st.session_state["final_case"] = ""
    if "retrieved_case" not in st.session_state:
        st.session_state["retrieved_case"] = ""
    if "retrieved_name" not in st.session_state:
        st.session_state["retrieved_name"] = ""
    if "selected_case_id" not in st.session_state:
        st.session_state["selected_case_id"] = -1
    if "lab_tests" not in st.session_state:
        st.session_state["lab_tests"] = {}
    if "selected_tests" not in st.session_state:
        st.session_state["selected_tests"] = {}
    if "lab_results" not in st.session_state:
        st.session_state["lab_results"] = ""
    if "selected_lab_result" not in st.session_state:
        st.session_state["selected_lab_result"] = None
    if "search_results" not in st.session_state:
        st.session_state["search_results"] = []
    if "learner_tasks" not in st.session_state:
        st.session_state["learner_tasks"] = learner_tasks

# Function to save a transcript
def save_transcript(transcript_content, role, specialty):
    new_transcript = Transcript(content=transcript_content, role=role, specialty=specialty)
    session.add(new_transcript)
    session.commit()

# Function to save an assessment
def save_assessment(assessment_content, role, specialty):
    new_assessment = Assessment(content=assessment_content, role=role, specialty=specialty)
    session.add(new_assessment)
    session.commit()

# Function to save case details
def save_case_details(case_details_content, saved_name, role = "", specialty=""):
    new_case_details = CaseDetails(content=case_details_content, saved_name=saved_name, role=role, specialty=specialty)
    session.add(new_case_details)
    session.commit()

# Function to save lab results
def save_lab_results(lab_results_content, saved_name):
    new_lab_results = LabResults(content=lab_results_content, saved_name=saved_name)
    session.add(new_lab_results)
    session.commit()

    # Debugging output
    saved_result = session.query(LabResults).filter_by(saved_name=saved_name).first()
    st.write(f"Debug: Saved Lab Result - {saved_result.saved_name}: {saved_result.content}")

# Function to retrieve records with full-text search and wildcards
def get_records(model, search_text=None, saved_name=None):
    query = session.query(model)
    if search_text:
        search_text = f"%{search_text}%"  # Wildcard search
        query = query.filter(model.content.ilike(search_text))
    if saved_name:
        saved_name = f"%{saved_name}%"  # Wildcard search
        query = query.filter(model.saved_name.ilike(saved_name))
    return query.all()

def llm_call(model, messages):
    try:
        response = requests.post(
            url="https://openrouter.ai/api/v1/chat/completions",
            headers={
                "Authorization": "Bearer " + st.secrets["OPENROUTER_API_KEY"],  # Fixed to correct access to secrets
                "Content-Type": "application/json",
                "HTTP-Referer": "https://fsm-gpt-med-ed.streamlit.app",  # To identify your app
                "X-Title": "lof-sims",
            },
            data=json.dumps({
                "model": model,
                "messages": messages,
            })
        )
    except requests.exceptions.RequestException as e:
        st.error(f"Error - make sure bills are paid!: {e}")
        return None
    # Extract the response content
    try:
        response_data = response.json()
    except json.JSONDecodeError:
        st.error(f"Error decoding JSON response: {response.content}")
        return None
    return response_data  # Adjusted to match expected JSON structure

def check_password():
    """Returns `True` if the user had the correct password."""

    def password_entered():
        """Checks whether a password entered by the user is correct."""
        if st.session_state["password"] == st.secrets["password"]:
            st.session_state["password_correct"] = True
            del st.session_state["password"]  # don't store password
        else:
            st.session_state["password_correct"] = False

    if "password_correct" not in st.session_state:
        # First run, show input for password.
        st.text_input(
            "Password", type="password", on_change=password_entered, key="password"
        )
        st.write("*Please contact David Liebovitz, MD if you need an updated password for access.*")
        return False
    elif not st.session_state["password_correct"]:
        # Password not correct, show input + error.
        st.text_input(
            "Password", type="password", on_change=password_entered, key="password"
        )
        st.error("😕 Password incorrect")
        return False
    else:
        # Password correct.
        return True

# Define the checklist fields
CHECKLIST_FIELDS = {
    "Onset": "",
    "Location": "",
    "Duration": "",
    "Character": "",
    "Aggravating/Alleviating factors": "",
    "Radiation": "",
    "Timing": "",
    "Severity": ""
}

def update_checklist_from_case(case_content):
    prompt = f"Extract relevant information for the following fields from the given case content: {', '.join(CHECKLIST_FIELDS.keys())}\n\nCase Content:\n{case_content}"
    response = llm_call("anthropic/claude-3-haiku", [{"role": "user", "content": prompt}])
    extracted_info = response['choices'][0]['message']['content'].strip().split('\n')
    
    checklist_fields = CHECKLIST_FIELDS.copy()
    for info in extracted_info:
        if ':' in info:
            field, value = info.split(':', 1)
            field = field.strip()
            value = value.strip()
            if field in checklist_fields:
                checklist_fields[field] = value
    return checklist_fields

def generate_checklist_template(checklist_fields):
    template = "HPI:\n\n"
    for field, value in checklist_fields.items():
        template += f"{field}: {value}\n\n"
    return template

st.title("Case Generator for Simulations")
init_session()

# Changes made by Tanaya
lab_tests_dict = {
    "Thyroid": {
        "TSH": "0.3-3.7 uIU/ml",
        "Free T3": "230-619 pg/dL",
        "Free T4": "0.7-1.9 ng/dL"
    },
    "Cardiac": {
        "CPK": "26-140 U/L",
        "CK-MB": "< 3% of total",
        "Troponin I": "0.00-0.04 ng/ml"
    },
    "CMP": {  # Abbreviation used here
        "Sodium": "135-145 mmol/L",
        "Potassium": "3.5-4.5 mmol/L",
        "Chloride": "98 - 106 mmol/L",
        "CO2": "22-29 mmol/L",
        "BUN": "7 - 18 mg/dL",
        "Cr": "0.6-1.2 mg/dL",
        "Glucose": "20-115 mg/dL",
        "Calcium": "8.4-10.2 mg/dL",
        "Magnesium": "1.3-2.1 mmol/L",
        "Alk Phos": "38-126 U/L",
        "Albumin": "3.5-5.5 g/dL",
        "Total Protein": "6 - 8 g/dL",
        "AST": "5 - 30 U/L",
        "ALT": "5 - 35 U/L",
        "Bilirubin": "< 1.0 mg/dL"
    },
    "CBC": {  # Abbreviation used here
        "WBC": "4,000 - 10,000 / mm3",
        "RBC": "4.6-6.2/microliter",
        "HGB": "13-16 g/dL",
        "HCT": "40-50%",
        "MCV": "80-100 fl",
        "MCH": "28-32 pg",
        "MCHC": "32-36 g/dL",
        "RDW": "11 - 14%",
        "Platelets": "140,000-150,000/mm3",
        "Neutrophils": "40 - 60%",
        "Lymphocytes": "20 - 40%",
        "Monocytes": "2 - 8%",
        "Eosinophils": "1 - 4%",
        "Basophils": "0.5% - 1%",
        "Band forms": "0 - 10%"
    },
    "UA": {  # Abbreviation used here
        "pH": "5 - 9",
        "Specific Gravity": "1.010 - 1.060",
        "Color": "Light yellow / colorless",
        "Turbidity": "Clear",
        "Protein": "NEG",
        "Glucose": "NEG",
        "Ketones": "NEG",
        "Bili": "NEG",
        "Blood": "NEG",
        "Leukocyte Esterase": "NEG",
        "Nitrite": "NEG",
        "WBC": "No Cells",
        "RBC": "No Cells",
        "Crystals": "None"
    },
    "ABG": {  # Abbreviation used here
        "pH": "7.34-7.44",
        "pCO2": "35-45 mmHg",
        "pO2": "75-100 mmHg",
        "HCO3": "22-26 mmol/L",
        "O2 Sat.": "95%-100%"
    },
    "Cholesterol": {
        "Total Cholesterol": "low risk <200, mod:200-239, high >239 mg/dL",
        "HDL": "at risk <40, optimal >59 mg/dL",
        "LDL": "low risk <130, mod:130-159, high >159 units/dL",
        "Triglycerides": "35-135(♀), 40-160(♂) mg/dL",
        "Chol/HDL Ratio": "1-3.5",
        "non-HDL Cholesterol": "<130 mg/dL"
    },
    "Coagulation": {  # Abbreviation used here
        "PT": "11-15 s",
        "aPTT": "20-35 s",
        "INR": "1",
        "D dimer": "<0.50 mcg/mL"
    },
    "Iron": {  # Abbreviation used here
        "serum iron": "37-145 (♀) 59-158 (♂) mcg/dL",
        "TIBC": "250-425 mcg/dL",
        "transferrin sat": "15-50% (♀) 20-50% (♂)",
        "ferritin": "12-150 (♀) 12-300 (♂) ng/mL"
    },
    
}
# Changes made by Tanaya - End

if check_password():
    st.info("Provide inputs and generate a case. After your case is generated, please click the '*Send case to the simulator!*' and then wake the simulator.")
    
    with st.expander("Model Options for Case Generation (Claude3 Haiku by default)", expanded=False):
        model_choice = st.selectbox("Model Options", (
            "anthropic/claude-3-haiku",
            "anthropic/claude-3-sonnet", 
            "anthropic/claude-3-opus", 
            "openai/gpt-4o", 
            "google/gemini-pro", 
            "meta-llama/llama-2-70b-chat",
        ), index=0)

    if "response_markdown" not in st.session_state:
        st.session_state["response_markdown"] = ""
        
    if "expanded" not in st.session_state:
        st.session_state["expanded"] = True
        
    if "edited_new_case" not in st.session_state:
        st.session_state["edited_new_case"] = ""
        
    if "learner_tasks" not in st.session_state:
        st.session_state["learner_tasks"] = learner_tasks
        
    tab1, tab2, tab3 = st.tabs(["New Case", "Retrieve a Case", "Lab Tests"])
    
    with tab1:
        # Move file upload section here
        st.header("Import a Suki File")

        uploaded_file = st.file_uploader("Choose a file", type=["txt", "md", "docx", "pdf"])
        
        if uploaded_file is not None:
            file_content = ""
            if uploaded_file.type == "text/plain":
                file_content = uploaded_file.read().decode("utf-8")
            elif uploaded_file.type == "text/markdown":
                file_content = uploaded_file.read().decode("utf-8")
            elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                doc = Document(uploaded_file)
                file_content = "\n".join([para.text for para in doc.paragraphs])
            elif uploaded_file.type == "application/pdf":
                pdf_document = fitz.open(stream=uploaded_file.read(), filetype="pdf")
                file_content = ""
                for page_num in range(pdf_document.page_count):
                    page = pdf_document.load_page(page_num)
                    file_content += page.get_text("text")

            st.session_state.final_case = file_content

            # Automatically update the checklist from the case content
            st.session_state.checklist_fields = update_checklist_from_case(file_content)
            st.session_state.checklist_template = generate_checklist_template(st.session_state.checklist_fields)

            if st.button("Submit Uploaded Case"):
                messages = [
                    {"role": "system", "content": f'Using the case details provided (supplemented as needed with additional generated content), comprehensively populate an educational clinical case description in the following specific format: {output_format}'},
                    {"role": "user", "content": file_content}
                ]
                with st.spinner("Assembling Case... Please wait."):
                    response_content = llm_call(model_choice, messages)
                if response_content:
                    st.session_state.response_markdown = response_content['choices'][0]['message']['content']
                    st.success("Case Study Framework Generated!")
                else:
                    st.error("Failed to generate case details. Please try again.")

        col1, col2, col3 = st.columns([2, 2, 5])

        with col1:    
            st.info("**Include desired history in the text paragraph. The AI will generate additional details as needed to draft an educational case.**")
                
            case_study_input = {
                'Case Title': st.text_input("Case Study Title", help="Presenting symptom, e.g."),
                'Case Description': st.text_area("Case Description", height=200, help = "As detailed or brief as desired, e.g., 65F with acute chest pain..."),
                'Case Primary Diagnosis': st.text_input("Primary Diagnosis", help = "The one or more primary diagnoses, e.g., Pulmonary Embolism"),
            }
            case_study_input = json.dumps(case_study_input)
            with st.expander("Default Learner Tasks", expanded = False):
                st.markdown(learner_tasks)  # Display the default tasks
            if st.checkbox("Edit Learner Tasks", value=False, key = "initial_case_edit"):
                learner_tasks = st.text_area("Learner Tasks for Assessment", height=200, help = "What the learner is expected to do, e.g., Perform a focused history and examination", value = learner_tasks)
            st.session_state.learner_tasks = learner_tasks
        
        with col1: 
            st.info("Click submit when ready to generate a case!")
            submit_button = st.button("Submit Case")

        if submit_button:
            messages = [
                {"role": "system", "content": f'Using the case details provided (supplemented as needed with additional generated content), comprehensively populate an educational clinical case description in the following specific format: {output_format}'},
                {"role": "user", "content": f'case_details: {case_study_input}'}
            ]
        
            with col2:
                with st.spinner("Assembling Case... Please wait."):
                    response_content = llm_call(model_choice, messages)
            st.session_state.response_markdown = response_content['choices'][0]['message']['content']
            st.success("Case Study Framework Generated!")

        with col3:
            if st.session_state.response_markdown != "" or st.session_state.final_case != "":
                with st.expander("View Full Case", expanded=st.session_state.expanded):
                    if st.session_state.response_markdown != "":
                        st.markdown(st.session_state.response_markdown)
                    else:
                        # Structure the content
                        content = st.session_state.final_case
                        lines = content.split('\n')
                        structured_content = ""
                        headers = ["Patient Approach", "History of Present Illness (HPI)", "Past Medical History (PMHx)", "Social History (SHx)", 
                                   "Family History (FHx)", "Medications and Allergies", "Review of Systems (ROS)", "Physical Examination", 
                                   "Diagnostic Reasoning", "Teaching Points"]
                        current_header = None

                        for line in lines:
                            if any(header in line for header in headers):
                                current_header = line
                                structured_content += f"### {line}\n"
                            else:
                                if current_header:
                                    structured_content += f"{line}\n"

                        st.markdown(structured_content)
                    st.session_state.expanded = False
            
                with col2:
                    st.info("Review and/or edit the case and begin the simulator!")                  
            
                    if st.checkbox("Edit Case (Scroll Down)", value=False, key="edit_case_checkbox"):
                        st.session_state.expanded = False
                        st.warning("Please edit the case as needed while leaving other characters, e.g., '#' and '*', in place. Use the 'Save Case Edits' button at the bottom to save edits!")
                        
                        edited_new_case = st.text_area("Click button at bottom to save your edits!", st.session_state.response_markdown if st.session_state.response_markdown != "" else st.session_state.final_case, height=1000) 
                        if st.button("Save Case Edits for the Simulator"):
                            st.success("Case Edits Saved!")
                            if edited_new_case:
                                st.session_state["final_case"] = edited_new_case
                        st.page_link("pages/🧠_Simulator.py", label="Wake the Simulator (including any saved edits)", icon="🧠")
                    else:
                        st.session_state["final_case"] = st.session_state.response_markdown if st.session_state.response_markdown != "" else st.session_state.final_case
                
                    if st.session_state.final_case !="":        
                        case_html = markdown2.markdown(st.session_state.final_case, extras=["tables"])
                            
                        if st.checkbox("Generate Case PDF file", key="generate_case_pdf"):
                            html_to_pdf(case_html, 'case.pdf')
                            with open("case.pdf", "rb") as f:
                                st.download_button("Download Case PDF", f, "case.pdf")
                        if st.checkbox("Generate Case DOCX file", key="generate_case_docx"):
                            html_to_docx(case_html, 'case.docx')
                            with open("case.docx", "rb") as f:
                                st.download_button("Download Case DOCX", f, "case.docx")

                    if st.session_state["final_case"] != "":
                        if st.button("Send case to the simulator!"):
                            st.session_state["final_case"] = st.session_state.final_case
                            st.session_state["retrieved_name"] = st.session_state.retrieved_name
                            st.session_state["learner_tasks"] = learner_tasks
                            st.page_link("pages/🧠_Simulator.py", label="Wake the Simulator", icon="🧠")

                st.divider()
                if st.checkbox("Save Case to the Database for Future Use"):
                    case_details = st.text_area("Case Details to Save to the Database for Future Use", value=st.session_state.final_case)
                    saved_name = st.text_input("Saved Name (Required to save case)")

                    if st.button("Save Case to the Database for future use!"):
                        if saved_name:
                            save_case_details(case_details, saved_name)
                            st.success("Case Details saved successfully!")
                        else:
                            st.error("Saved Name is required to save the case")

    if "search_results" not in st.session_state:
        st.session_state.search_results = []
    if "selected_case" not in st.session_state:
        st.session_state.selected_case = None

    # Tab2 content for retrieving and selecting cases
    with tab2:
        
        col3, col4 = st.columns([1,3])
        with col3:
            st.header("Retrieve Records")
            search_text = st.text_input("Search Text")
            search_saved_name = st.text_input("Search by Saved Name")
            search_role =""
            search_specialty = ""
            if search_role in ["Resident", "Fellow", "Attending"]:
                search_specialty = st.text_input("Search by Specialty", "")

            if st.button("Search Cases"):
                st.session_state.search_results = get_records(CaseDetails, search_text, search_saved_name)

            if st.session_state.search_results:
                st.subheader("Cases Found")
                for i, case in enumerate(st.session_state.search_results, start=1):
                    st.write(f"{i}. {case.saved_name}")
                    if st.button(f"View (and Select) Case {i}", key=f"select_case_{i}"):
                        st.session_state.selected_case = case
            with col4:
                if st.session_state.selected_case:
                    st.subheader("Retrieved Case")
                    with st.expander("View Full Case", expanded=False):
                        st.write(f'Here is the retrieved case name: {st.session_state.selected_case.saved_name}')
                        st.write(st.session_state.selected_case.content)
                        st.session_state.final_case = st.session_state.selected_case.content
                    if st.checkbox("Edit Retrieved Case (Scroll Down)", value=False, key="retrieved_case_edit"):
                        st.session_state.expanded = False
                        st.warning('Please edit the case as needed while leaving other characters, e.g., "#" and "*", in place. Remember to update the Door Chart section at the bottom!')
                        updated_retrieved_case = st.text_area("Edit Case, enter control-enter or command-enter to save edits!", st.session_state.selected_case.content, height=1000)
                        make_new_entry = st.checkbox("If desired, make a database entry when saving edits.", value=False, key="make_new_entry_checkbox")
                        if make_new_entry:
                            saved_name = st.text_input("Saved Name after edits (Required to save case)")
                        if st.button("Save Edits for Simulation and Generate a PDF"):
                            st.session_state.final_case = updated_retrieved_case
                            st.info("Case Edits Saved!")   
                            updated_case_html = markdown2.markdown(st.session_state.final_case, extras=["tables"])
                            html_to_pdf(updated_case_html, 'updated_case.pdf')
                            with open("updated_case.pdf", "rb") as f:
                                st.download_button("Download Updated Case PDF", f, "updated_case.pdf")
                            if st.button("Generate Case DOCX file"):
                                html_to_docx(updated_case_html, 'updated_case.docx')
                                with open("updated_case.docx", "rb") as f:
                                    st.download_button("Download Case DOCX", f, "updated_case.docx")
                            if make_new_entry:
                                if saved_name:
                                    save_case_details(st.session_state.final_case, saved_name)
                                    st.success("Case Details saved successfully!")
                                else:
                                    st.error("Saved Name is required to save the case")
                    st.page_link("pages/🧠_Simulator.py", label="Wake the Simulator ", icon="🧠")

   # Changes made by Tanaya
    # Tab3
    with tab3:
        st.header("Generate Lab Test Results")

        primary_diagnosis = st.text_input("Primary Diagnosis")

        col1, col2 = st.columns([1, 1])
        
        with col1:
            selected_diagnoses = st.multiselect("Select Panel for Lab Tests", options=list(lab_tests_dict.keys()))
            if selected_diagnoses:
                st.session_state.lab_tests = {diagnosis: lab_tests_dict[diagnosis] for diagnosis in selected_diagnoses}
                st.session_state.selected_tests = {diagnosis: list(lab_tests_dict[diagnosis].keys()) for diagnosis in selected_diagnoses}

                selected_tests_data = []
                for diagnosis, tests in st.session_state.lab_tests.items():
                    for test, normal_range in tests.items():
                        selected_tests_data.append([diagnosis, test, normal_range])

                df = pd.DataFrame(selected_tests_data, columns=["Panel", "Test", "Normal Range"])
                df["Select"] = True  # Default to all selected
                edited_df = st.data_editor(df, use_container_width=True, height=200)  # Adjust height

                # Filter selected tests based on user input
                selected_tests = {}
                for i, row in edited_df.iterrows():
                    if row['Select']:
                        diagnosis = row['Panel']
                        test = row['Test']
                        if diagnosis not in selected_tests:
                            selected_tests[diagnosis] = []
                        selected_tests[diagnosis].append(test)
                st.session_state.selected_tests = selected_tests

        with col2:
            prompt = st.text_area("Specify details for LLM to generate lab results with normality or abnormality", height=232, key="global_prompt")  # Adjust height

        col_full = st.columns([1])
        with col_full[0]:
            if st.button("Generate Lab Test Results"):
                if primary_diagnosis:
                    all_lab_results = ""
                    for diagnosis, tests in st.session_state.selected_tests.items():
                        if tests:
                            messages = [
                                {"role": "system", "content": "You are an AI that generates realistic lab test results based on the provided primary diagnosis and selected tests."},
                                {"role": "user", "content": f"Primary Diagnosis: {primary_diagnosis}\nSelected Tests: {json.dumps(tests)}\nDetails: {prompt}"}
                            ]
                            with st.spinner(f"Generating lab test results for {diagnosis}..."):
                                response_content = llm_call(model_choice, messages)
                            if response_content:
                                lab_results = response_content['choices'][0]['message']['content']
                                all_lab_results += f"### Lab Results for {diagnosis}\n{lab_results}\n\n"
                            else:
                                st.error(f"Failed to generate lab test results for {diagnosis}.")
                    
                    st.session_state.lab_results = all_lab_results
                    with st.expander("Lab Results", expanded=True):
                        st.markdown(st.session_state.lab_results, unsafe_allow_html=True)
                    lab_results_html = markdown2.markdown(st.session_state.lab_results, extras=["tables"])
                    
                    # Download buttons for PDF and DOCX
                    html_to_pdf(lab_results_html, 'lab_results.pdf')
                    with open("lab_results.pdf", "rb") as f:
                        st.download_button("Download Lab Results PDF", f, "lab_results.pdf")
                    
                    html_to_docx(lab_results_html, 'lab_results.docx')
                    with open("lab_results.docx", "rb") as f:
                        st.download_button("Download Lab Results DOCX", f, "lab_results.docx")
