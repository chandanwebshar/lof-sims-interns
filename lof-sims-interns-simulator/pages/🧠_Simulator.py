import streamlit as st
from sim_prompts import *
import markdown2
from groq import Groq
from openai import OpenAI
import os
from bs4 import BeautifulSoup
#from fpdf import FPDF
from datetime import datetime
from audio_recorder_streamlit import audio_recorder
from prompts import *
import tempfile
import requests
import json
import base64
import random
from Start import llm_call
from docx import Document as DocxDocument
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import requests
import streamlit as st
import base64
from openai import OpenAI
import uuid
import asyncio
from aiortc import RTCPeerConnection, RTCSessionDescription
from aiortc.exceptions import OperationError

import requests
from aiortc import RTCPeerConnection, RTCSessionDescription

# st.set_page_config(page_title='Simulated Chat', layout = 'centered', page_icon = ':stethoscope:', initial_sidebar_state = 'expanded')


def assign_random_voice(sex):
    """
    Randomly assigns one of the specified strings to the variable 'voice'.

    Returns:
    - str: The assigned voice.
    
    The possible voices are 'alloy', 'echo', 'fable', 'onyx', 'nova', and 'shimmer'.
    """
    # List of possible voices
    male_voices = [ 'echo', 'fable', 'onyx' ]
    female_voices = [ 'nova', 'shimmer']
    
    if sex == 'male':
        voices = male_voices
    else:
        voices = female_voices
    
    # Randomly voice one voice from the list
    voice = random.choice(voices)
    
    return voice

class Word():
    def __init__(self, title):
       self.document = DocxDocument()
       self.title = title
       self.add_title()
    def add_title(self):
        self.document.add_heading(self.title, level=1)

    def header(self):
        self.set_font("Arial", "B", 12)
        self.cell(0, 10, self.title, 0, 1, "C")
        #self.ln(10)
        

    def chapter_title(self, title, level=1):
        heading = self.document.add_heading(level=level)
        run = heading.add_run(title)
        run.font.size = Pt(16 if level == 1 else 14 if level == 2 else 12)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    def chapter_body(self, body):
        #self.set_font("Arial", "", 12)
        #self.multi_cell(0, 10, body)
        #self.ln()
        #run.font.size = Pt(12)
        
        paragraph = self.document.add_paragraph(body)
        for run in paragraph.runs:
            run.font.size = Pt(12)
    def add_page(self):
        self.document.add_page_break()

    #def add_page_break(self):
    #    self.add_page_break()

    def add_list(self, items, is_ordered=False):
        #self.set_font("Arial", "", 12)
        for i, item in enumerate(items, start=1):
            if is_ordered:
                paragraph = self.document.add_paragraph(f"{i}. {item}")
            else:
                paragraph = self.document.add_paragraph(f"- {item}")
            run = paragraph.runs[0]
            run.font.size = Pt(12)
    def output(self,name):
        self.document.save(name)

def transcript_to_doc(html_content, name):   
     # Use BeautifulSoup to parse the HTML
    html_content = html_content.replace('🤒', 'Patient').replace('👩⚕️', 'Doctor')
    html_content = html_content.encode('latin-1', 'ignore').decode('latin-1')
    soup = BeautifulSoup(html_content, "html.parser")
    
    
    # Extract title for the document
    title = "Patient Case"
    
    # Create PDF instance and set the title
    word = Word(title)
    word.title = title
    #doc.add_page_break()
    #doc.set_auto_page_break(auto=True, margin=15)
    #doc.set_font("Arial", size=12)
    

    # Process each section of the HTML
    for element in soup.find_all(["h2", "h3", "p", "ul", "ol", "li", "hr"]):
        if element.name == "h2":
            word.chapter_title(element.get_text(), level=2)
        elif element.name == "h3":
            word.chapter_title(element.get_text(), level=3)
        elif element.name == "p":
            word.chapter_body(element.get_text())
        elif element.name == "ul":
            items = [li.get_text() for li in element.find_all("li")]
            word.add_list(items, is_ordered=False)
        elif element.name == "ol":
            items = [li.get_text() for li in element.find_all("li")]
            word.add_list(items, is_ordered=True)
        elif element.name == "hr":
            word.add_page()
    
    # Output the PDF
    word.output(name)

def llm_call(model, messages):
    
    try:
        response = requests.post(
            url="https://openrouter.ai/api/v1/chat/completions",
            headers={
                "Authorization": "Bearer " + st.secrets["OPENROUTER_API_KEY"],  # Fixed to correct access to secrets
                "Content-Type": "application/json",
                "HTTP-Referer": "https://fsm-gpt-med-ed.streamlit.app",  # To identify your app
                "X-Title": "lof-sims",
            },
            data=json.dumps({
                "model": model,
                "messages": messages,
            })
        )
    except requests.exceptions.RequestException as e:
        st.error(f"Error - make sure bills are paid!: {e}")
        return None
    # Extract the response content
    try:
        response_data = response.json()
    except json.JSONDecodeError:
        st.error(f"Error decoding JSON response: {response.content}")
        return None
    return response_data  # Adjusted to match expected JSON structure


def parse_groq_stream(stream):
    for chunk in stream:
        if chunk.choices:
            if chunk.choices[0].delta.content is not None:
                yield chunk.choices[0].delta.content

base_url = "https://api.d-id.com"


def get_agents():

    url = "https://api.d-id.com/agents/me?limit=100"

    headers = {
        "accept": "application/json",
        "authorization": "Basic c2hpa2hhc2hhcm1hOTYzN0BnbWFpbC5jb20:SeojuJBCjwo2aEbOHtLF7"
    }

    response = requests.get(url, headers=headers)

    print(response.text)


def create_stream():

    url = "https://api.d-id.com/talks/streams"

    payload = {
        "stream_warmup": "false",
        "source_url": "https://create-images-results.d-id.com/DefaultPresenters/Noelle_f/v1_image.jpeg"
    }
    headers = {
        "accept": "application/json",
        "content-type": "application/json",
        "authorization": "Basic YzJocGEyaGhjMmhoY20xaE9UWXpOMEJuYldGcGJDNWpiMjA6U2VvanVKQkNqd28yYUViT0h0TEY3Og=="
    }

    response = requests.post(url, json=payload, headers=headers)
    #return response.json()
    #print(response.text)
    
    return response

def transcribe_audio(audio_file_path):
    from openai import OpenAI
    api_key = st.secrets["OPENAI_API_KEY"]
    client = OpenAI(    
        base_url="https://api.openai.com/v1",
        api_key=api_key,
    )
    audio_file = open(audio_file_path, "rb")
    transcript = client.audio.transcriptions.create(
    model="whisper-1", 
    file=audio_file, 
    response_format="text"
    )
    return transcript

def talk_stream(model, voice, input):
    api_key = st.secrets["OPENAI_API_KEY"]
    client = OpenAI(    
        base_url="https://api.openai.com/v1",
        api_key=api_key,
    )
    try:
        response = client.audio.speech.create(
        model= model,
        voice= voice,
        input= input,
        )
        response.stream_to_file("last_interviewer.mp3")
    
    except Exception as e:
        st.write("The API is busy - should work in a moment for voice.")



    
def autoplay_local_audio(filepath: str):
    # Read the audio file from the local file system
    with open(filepath, 'rb') as f:
        data = f.read()
    b64 = base64.b64encode(data).decode()
    md = f"""
        <audio controls autoplay="true">
        <source src="data:audio/mp3;base64,{b64}" type="audio/mp3">
        </audio>
        """
    st.markdown(
        md,
        unsafe_allow_html=True,
    )


@st.cache_data
def extract_patient_door_chart_section(text):
    """
    Extracts the PATIENT DOOR CHART section from the given text string and returns it.
    
    Args:
    - text (str): The input text containing multiple sections, including "PATIENT DOOR CHART".
    
    Returns:
    - str: The extracted "PATIENT DOOR CHART" section through the end of the provided text.
    """
    # Define the start marker for the section to extract
    start_marker = "## PATIENT DOOR CHART"
    
    # Find the position where the relevant section starts
    start_index = text.find(start_marker)
    
    # If the section is found, extract and return the text from that point onwards
    if start_index != -1:
        return text[start_index:]
    else:
        # Return a message indicating the section was not found if it doesn't exist in the string
        return "PATIENT DOOR CHART section not found in the provided text. Please go back to the Start page!"
# st.write(f'Here is the case {st.session_state.final_case}')

try:
    extracted_section = extract_patient_door_chart_section(st.session_state.final_case)
    st.info(extracted_section)
    st.info(st.session_state.learner_tasks)
    if "messages" not in st.session_state:
        st.session_state.messages = [{"role": "system", "content": f'{sim_persona} Here are the specifics for your persona: {st.session_state.final_case}'}, ]
except Exception as e:
    st.error(f"Please return to the main page. An error occurred. Please do not 're-load' when in the middle of a conversation. Here are the error details: {e}. ")

if "sex" not in st.session_state:
    st.session_state.sex = ""
if st.session_state.sex == "":
    messages_sex =[{"role": "user", "content": f'Analyze the following content and return only the sex, e.g., male, female, or other. Return nothing else. {extracted_section}'}]
    st.session_state.sex = llm_call("anthropic/claude-3-haiku", messages_sex)

#################################################

# Set OpenAI API key from Streamlit secrets
groq_client = Groq(api_key = st.secrets['GROQ_API_KEY'])

# st.set_page_config(
#     page_title='Fast Helpful Chat',
#     page_icon='🌌',
#     initial_sidebar_state='expanded'
# )


st.title("Clinical Simulator Chat")
# st.caption('Powered by [Groq](https://groq.com/).')

if "password_correct" not in st.session_state:
    st.session_state["password_correct"] = False
    st.write("Login on the Sims page to get started.")

if st.session_state["password_correct"] == True:
    st.info("Type your questions at the bottom of the page or use voice input (left sidebar)! You may need to right click your Chrome browser tab to unmute this website and also accept the microphone permissions.")

    
    # st.sidebar.title('Customization')
    with st.sidebar:
        with st.expander("Change Model", expanded=False):
            st.session_state.model = st.selectbox(
                    'voice a model',
                    ['llama3-70b-8192', 'gpt-4o',], index=1,
                )
        # Initialize chat history

        
    # if st.sidebar.checkbox("Change personality? (Will clear history.)"):
    #     persona = st.sidebar.radio("Pick the persona", ("Regular user", "Physician"), index=1)
    #     if persona == "Regular user":
    #         system = st.sidebar.text_area("Make your own system prompt or use as is:", value=system_prompt2)
    #     else:
    #         system = system_prompt
    #     st.session_state.messages = [{"role": "system", "content": system}]
        
    if "sim_response" not in st.session_state:
        st.session_state["sim_response"] = ""

    if "audio_off" not in st.session_state:
        st.session_state["audio_off"] = False

    if "audio_input" not in st.session_state:
        st.session_state["audio_input"] = ""
        
    if "voice" not in st.session_state:
        # Example usage:
        st.session_state["voice"] = assign_random_voice(st.session_state.sex)
        
    if "results" not in st.session_state:
        st.session_state["results"] = ""
        
    if "orders_placed" not in st.session_state:
        st.session_state["orders_placed"] = ""
        
    if "conversation_string" not in st.session_state:
        st.session_state["conversation_string"] = ""
        
    if "assessment" not in st.session_state:
        st.session_state["assessment"] = ""
        
    if "last_audio_size" not in st.session_state:
        st.session_state["last_audio_size"] = 0
        
    if "h_and_p" not in st.session_state:
        st.session_state["h_and_p"] = ""
    if "suggested_orders" not in st.session_state:
        st.session_state["suggested_orders"] = ""
    

            # Audio selection
    
    input_source = st.sidebar.radio("Choose to type or speak!", ("Text", "Microphone"), index=0)
    st.session_state.audio_off = st.sidebar.checkbox("Turn off voice response", value=False) 
    # Display chat messages from history on app rerun
    conversation_str = extracted_section + "**Learner Tasks:**\n\n" + st.session_state.learner_tasks + "\n\n" + "______" + "\n\n" + "**Clinical Interview:**\n\n"
    for message in st.session_state.messages:
        if message["role"] == "user":
            with st.chat_message(message["role"], avatar="👩⚕️"):
                st.markdown(message["content"])
                conversation_str += "👩⚕️: " + message["content"] + "\n\n"
        elif message["role"] == "assistant":
            with st.chat_message(message["role"], avatar="🤒"):
                st.markdown(message["content"])
                conversation_str += "🤒: " + message["content"] + "\n\n"
    conversation_str += "______" + "\n\n" + "**Orders:**\n\n" + st.session_state.orders_placed +  "**Results:**\n\n""\n\n" + st.session_state.results + "\n\n"
    st.session_state.conversation_string = conversation_str



    if input_source == "Text":
    
    # Accept user input
        if prompt := st.chat_input("What's up?"):
            # Add user message to chat history
            st.session_state.messages.append({"role": "user", "content": prompt})
            # Display user message in chat message container
            with st.chat_message("user", avatar="👩⚕️"):
                st.markdown(prompt)
                
                # Display assistant response in chat message container
            with st.chat_message("assistant", avatar="🤒"):    
                if st.session_state.model == "llama3-70b-8192":    
                    stream = groq_client.chat.completions.create(
                        model=st.session_state["model"],
                        messages=[
                            {"role": m["role"], "content": m["content"]}
                            for m in st.session_state.messages
                        ],
                        temperature=0.3,
                        stream=True,
                    )
                    st.session_state.sim_response = st.write_stream(parse_groq_stream(stream))
                    
                elif st.session_state.model == "gpt-4o":
                    api_key = st.secrets["OPENAI_API_KEY"]
                    client = OpenAI(
                            base_url="https://api.openai.com/v1",
                            api_key=api_key,
                    )
                    completion = client.chat.completions.create(
                        model = st.session_state.model,
                        messages = [
                            {"role": m["role"], "content": m["content"]}
                            for m in st.session_state.messages
                        ],
                        # headers={ "HTTP-Referer": "https://fsm-gpt-med-ed.streamlit.app", # To identify your app
                        #     "X-Title": "GPT and Med Ed"},
                        temperature = 0.5,
                        max_tokens = 1000,
                        stream = True,   
                        )     
                
                    # placeholder = st.empty()
                    st.session_state.sim_response = st.write_stream(completion)
                    
                    
                
            st.session_state.messages.append({"role": "assistant", "content": st.session_state.sim_response})
    else:
        with st.sidebar:
            st.info("Click the green person-icon, pause 3 seconds, and begin to speak with natural speech.\
                    As soon as you pause, the LLM will start its response.")
            audio_bytes = audio_recorder(
                text="Click, pause, speak:",
                recording_color="#e8b62c",
                neutral_color="#6aa36f",
                icon_name="user",
                icon_size="3x",
            )

        if audio_bytes:
            try:
                # Save audio bytes to a temporary file
                with tempfile.NamedTemporaryFile(delete=False, suffix=".wav") as temp_file:
                    temp_file.write(audio_bytes)
                    audio_file_path = temp_file.name
                    
                                # Inform about audio file size
                file_stats = os.stat(audio_file_path)
                # st.write("We have audio bytes!!! Length: ", file_stats.st_size)
                if st.session_state["last_audio_size"] != file_stats.st_size:

                    with st.spinner("Transcribing audio... Please wait."):
                        prompt = transcribe_audio(audio_file_path)

                    st.session_state.messages.append({"role": "user", "content": prompt})

                    # Display user message in chat message container
                    with st.chat_message("user", avatar="👩⚕️"):
                        st.markdown(prompt)
                        
                    



            finally:
                # Ensure the tempfile is removed regardless of success or failure in processing
                if 'audio_file_path' in locals():
                    os.remove(audio_file_path)
                    # st.write("Temporary audio file removed.")

            # Clearing audio bytes manually, might be redundant if no other operations store this variable
            audio_bytes = None

            if st.session_state["last_audio_size"] != file_stats.st_size:    
                # Display assistant response in chat message container
                with st.chat_message("assistant", avatar="🤒"):
                    with st.spinner("Answering... Please wait."):     
                        if st.session_state.model == "llama3-70b-8192":   
                            stream = groq_client.chat.completions.create(
                                model=st.session_state["model"],
                                messages=[
                                    {"role": m["role"], "content": m["content"]}
                                    for m in st.session_state.messages
                                ],
                                temperature=0.3,
                                stream=True,
                            )
                            st.session_state.sim_response = st.write_stream(parse_groq_stream(stream))
                        elif st.session_state.model == "gpt-4o":
                            api_key = st.secrets["OPENAI_API_KEY"]
                            client = OpenAI(
                                    base_url="https://api.openai.com/v1",
                                    api_key=api_key,
                            )
                            completion = client.chat.completions.create(
                                model = st.session_state.model,
                                messages = [
                                    {"role": m["role"], "content": m["content"]}
                                    for m in st.session_state.messages
                                ],
                                # headers={ "HTTP-Referer": "https://fsm-gpt-med-ed.streamlit.app", # To identify your app
                                #     "X-Title": "GPT and Med Ed"},
                                temperature = 0.5,
                                max_tokens = 1000,
                                stream = True,   
                                )     
                        
                            # placeholder = st.empty()
                            st.session_state.sim_response = st.write_stream(completion)
                        
                        
                    
                st.session_state.messages.append({"role": "assistant", "content": st.session_state.sim_response})
                st.session_state["last_audio_size"] = file_stats.st_size
                
    
    if st.session_state.audio_off == False:

        if st.session_state.sim_response:
            with st.spinner("Synthesizing audio... Please wait."):
                talk_stream("tts-1", st.session_state.voice, st.session_state.sim_response)
            autoplay_local_audio("last_interviewer.mp3")
            st.info("Note - this is an AI synthesized voice.")            
            st.session_state.sim_response = "" 
            os.remove("last_interviewer.mp3")   
                

    # if st.session_state["sim_response"]:
    #     conversation_str = ""
    #     for message in st.session_state.messages:
    #         if message["role"] == "user":
    #             conversation_str += "👩⚕️: " + message["content"] + "\n\n"
    #         elif message["role"] == "assistant":
    #             conversation_str += "🤒: " + message["content"] + "\n\n"
    #     st.session_state.conversation_string = conversation_str
    #     html = markdown2.markdown(conversation_str, extras=["tables"])
    #     st.download_button('Download the conversation when done!', html, f'sim_response.html', 'text/html')
    #     st.session_state.sim_response = ""
    
    st.sidebar.divider()
    st.sidebar.subheader("Chart Access")
    

    
    orders = st.sidebar.checkbox("Place Orders/Take Actions", value=False)
    if orders:
        with st.sidebar:
            # Ensure the 'suggestions_text' key exists in the session state
            if "suggestions_text" not in st.session_state:
                # Generating suggestions for the placeholder
                placeholder_prompt = f"Based on the following case details, list in one word the various medical tests to place, medications to take and lab tests to take:\n\n{st.session_state.final_case}"
                placeholder_messages = [{"role": "user", "content": placeholder_prompt}]
                
                with st.spinner("Generating suggestions..."):
                    placeholder_results = llm_call("openai/gpt-4o", placeholder_messages)
                
                # Get the suggestions to display in the expander
                if placeholder_results:
                    st.session_state.suggestions_text = placeholder_results['choices'][0]['message']['content']
                else:
                    st.session_state.suggestions_text = "No suggestions available at the moment."

            # Text input for order details without suggestions as placeholder
            order_details = st.text_input("Order Details", value="", placeholder="E.g., examine lungs, CXR, CBC, furosemide 40 mg IV x 1, consult cardiology, etc.", key="order")

            # Expander to show the suggestions
            with st.expander("Suggestions for Orders/Actions", expanded=False):
                suggestions = st.session_state.suggestions_text.split('\n\n')
                medical_tests = []
                medications = []
                lab_tests = []

                # Classify suggestions into categories
                for suggestion in suggestions:
                    if "medical" in suggestion.lower():
                        medical_tests.append(suggestion.strip())
                    elif "medication" in suggestion.lower() or "med" in suggestion.lower():
                        medications.append(suggestion.strip())
                    elif "lab" in suggestion.lower():
                        lab_tests.append(suggestion.strip())

                checked_suggestions = {
                    "Medical Tests": [],
                    "Medications": [],
                    "Lab Tests": []
                }

                st.subheader("Medical Tests")
                for i, suggestion in enumerate(medical_tests):
                    if suggestion:
                        if st.checkbox(suggestion, key=f"medical_test_{i}"):
                            checked_suggestions["Medical Tests"].append(suggestion)
                
                st.subheader("Medications")
                for i, suggestion in enumerate(medications):
                    if suggestion:
                        if st.checkbox(suggestion, key=f"medication_{i}"):
                            checked_suggestions["Medications"].append(suggestion)
                
                st.subheader("Lab Tests")
                for i, suggestion in enumerate(lab_tests):
                    if suggestion:
                        if st.checkbox(suggestion, key=f"lab_test_{i}"):
                            checked_suggestions["Lab Tests"].append(suggestion)

            if st.button("Submit Orders/Take Actions"):
                # Get the current date and time
                current_datetime = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                order_details_with_datetime = f"{order_details}\n\nDate and Time of Request: {current_datetime}"
                
                # Update session state with the new orders
                combined_suggestions = ""
                for category, suggestions_list in checked_suggestions.items():
                    if suggestions_list:  # Only add category if there are checked suggestions
                        combined_suggestions += "\n".join(suggestions_list) + "\n\n"
            
                st.session_state.orders_placed = order_details_with_datetime + "\n\n" + combined_suggestions

                # Prompt for orders/actions based on the case details and current orders
                prompt = orders_prompt.format(order_details=order_details, case_details=st.session_state.final_case, order_datetime=current_datetime, prior_results=st.session_state.results)
                orders_messages = [{"role": "user", "content": prompt}]
                with st.spinner("Transmitting Orders... Please wait."):
                    orders_results = llm_call("openai/gpt-4o", orders_messages)
                    
                st.session_state.results = orders_results['choices'][0]['message']['content'] + "\n\n" + st.session_state.results
                
            with st.expander("Completed Orders/Actions", expanded=False):
                st.write(st.session_state.orders_placed)
                
            with st.expander("All Results of Orders/Actions", expanded=False):
                st.write(st.session_state.results)




    html2 = markdown2.markdown(st.session_state.conversation_string, extras=["tables"])
    # st.sidebar.download_button('Download the transcript!', html2, f'transcript.html', 'text/html')
    
    with st.sidebar:

        h_and_p = st.checkbox("Generate a History and Physical (no A/P section)", value=False)
        if h_and_p:
            prompt = h_and_p_prompt.format(conversation_transcript=st.session_state.conversation_string)
            h_and_p_messages = [{"role": "user", "content": prompt}]
            if st.sidebar.button("Create the History and Physical"):
                with st.sidebar:
                    with st.spinner("Writing History and Physical... Please wait."):
                        try:
                            h_and_p_response = llm_call("anthropic/claude-3-sonnet", h_and_p_messages)
                        except Exception as e:
                            st.error("Error formulating history and physical. Here are the error details: " + str(e))
                st.session_state.h_and_p = h_and_p_response['choices'][0]['message']['content']
                
            if st.session_state.h_and_p:
                with st.expander("History and Physical", expanded = False):
                    st.write(st.session_state.h_and_p)
                html = markdown2.markdown(st.session_state.h_and_p, extras=["tables"])
                # st.sidebar.download_button('Download the assessment when done!', html, f'assessment.html', 'text/html')
                with st.sidebar:
                    if st.button("Generate H&P Word file"):
                        transcript_to_doc(html, 'h_and_p.docx')
                        with open("h_and_p.docx", "rb") as f:
                            st.download_button("Download H&P Word", f, "h_and_p.docx")
        
        
        
        
        st.divider()     
        if st.button("Generate Transcript Doc file"):
            transcript_to_doc(html2, 'transcript.docx')
            with open("transcript.docx", "rb") as f:
                st.download_button("Download Transcript Doc", f, "transcript.docx")    
        assess = st.checkbox("Assess Interaction", value=False)
    
    if assess:
        student_level = st.sidebar.selectbox("Student Level", ["1st Year Medical Student", "2nd Year Medical Student", "3rd Year Medical Student", "4th Year Medical Student"])
        prompt = assessment_prompt.format(learner_tasks = st.session_state.learner_tasks, student_level = student_level, case_details=st.session_state.final_case, conversation_transcript=st.session_state.conversation_string, orders_placed=st.session_state.orders_placed, results=st.session_state.results)
        assessment_messages = [{"role": "user", "content": prompt}]
        if st.sidebar.button("Formulate Assessment"):
            with st.sidebar:
                with st.spinner("Formulating Assessment... Please wait."):
                    try:
                        assessment_response = llm_call("anthropic/claude-3-sonnet", assessment_messages)
                    except Exception as e:
                        st.error("Error formulating assessment, be sure to download the transcript and try again. Here are the error details: " + str(e))
            st.session_state.assessment = assessment_response['choices'][0]['message']['content']
        
        if st.session_state.assessment:
            with st.expander("Assessment", expanded = False):
                st.write(st.session_state.assessment)
            html = markdown2.markdown(st.session_state.assessment, extras=["tables"])
            # st.sidebar.download_button('Download the assessment when done!', html, f'assessment.html', 'text/html')
            with st.sidebar:
                if st.button("Generate Assessment Doc file"):
                    transcript_to_doc(html, 'assessment.docx')
                    with open("assessment.docx", "rb") as f:
                        st.download_button("Download Assessment Doc", f, "assessment.docx")
                # st.divider()         
                # assess = st.checkbox("Assess Interaction", value=False)
        
        
